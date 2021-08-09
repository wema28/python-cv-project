from docx import Document
from docx.shared import Inches
import docx
import pyttsx3 

def speak(text):
    pyttsx3.speak(text)


document = Document()

# adding picture
document.add_picture(
    "pexels-jimmy-jimmy-1484794.jpg", 
    width = Inches(2.0)
)

# name, phone number and email
name = input("What is your name? ")
phone_number = input("What is your phone number? ")
email = input("What is your email? ")

# adding text
document.add_paragraph(
    name + " | " + phone_number + " | " + email)

# about me
document.add_heading("About Me")
about_me = input("Tell me something about you: ")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work experience ")
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From Date: ")
to_date = input("To Date: ")

p.add_run(company + ": ").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True

experience_details = input("The experience at " + company + " was: ")
profession = input("What was your profession at " + company + "? ")

p.add_run("The experience was " + experience_details + "\n")
p.add_run("I was a " + profession + " at " +company + "\n")

# more experiences with while loop
while True:
    has_more_experiences = input(
        "Do you have more experiences? Yes or No? ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From Date: ")
        to_date = input("To Date: ")

        p.add_run(company + ": ").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True

        experience_details = input("The experience at " + company + " was: ")
        profession = input("What was your profession at " + company + "? ")

        p.add_run("The experience was " + experience_details + "\n")
        p.add_run("I was a " + profession + " at " +company + "\n")

    else:
        break

# skills
document.add_heading("Skills")
skills = input("What is your best Skill? ")
skills = input("What is your other Skill? ")
document.add_paragraph(skills, style='List Bullet')

while True:
    has_more_skills = input("Do you have more Skills? Yes or No? ")
    if has_more_skills.lower() == "yes":
        skills = input("What Skill do you also have? ")
        document.add_paragraph(skills, style='List Bullet')
    
    else:
        break


document.save("cv.docx")